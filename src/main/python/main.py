from fbs_runtime.application_context.PyQt5 import ApplicationContext
from PyQt5.QtGui import QIntValidator
from PyQt5.QtWidgets import *
from PyQt5 import QtGui, QtCore
from functools import partial
from pathlib import Path
from os import path
import os
import docx
import sys


# Internal imports
import random

# External imports
import nltk
from nltk.corpus import wordnet


def download():
    nltk.download('popular')

def synoantonym_string(string_input, excluded_words, synonym):
    similar = synonym
    words = string_input.split(' ')

    for i in range(len(words)):
        word = words[i]

        special = ".?!\",(){}[]\\/-~\n"
        start_special = ""
        end_special = ""
        for char in word:
            if char in special:
                start_special = start_special + char
            else:
                break

        for char in word[::-1]:
            if char in special:
                end_special = end_special + char
            else:
                break

        word = word.replace(start_special, "", 1)
        if not end_special == '':
            word = "".join(word.rsplit(end_special, 1))

        if not word.lower() in excluded_words:

            new_val = ""
            synonyms = []
            antonyms = []

            for syn in wordnet.synsets(word):
                for l in syn.lemmas():
                    synonyms.append(l.name())
                    if l.antonyms():
                        antonyms.append(l.antonyms()[0].name())

            synonyms = set(synonyms)
            antonyms = set(antonyms)

            syn_len = len(synonyms)
            ant_len = len(antonyms)

            if similar:
                if syn_len == 0:
                    new_val = word
                else:
                    new_val = list(synonyms)[random.randint(0, syn_len - 1)]
            else:
                if ant_len == 0:
                    new_val = word
                else:
                    new_val = list(antonyms)[random.randint(0, ant_len - 1)]

            new_val = new_val.replace("_", " ")
            new_val = start_special + new_val + end_special

            words[i] = new_val

    string_output = " ".join(words)

    return string_output


def synoantonym_file(file_name, excluded_words, synonym):
    # It kept trying to replace it with IT or information technology lol

    with open(file_name) as input_file, open(file_name[:-4] + "_output.txt", "w+") as output_file:
        for line in input_file:
            line = line.strip(" \n")
            if not line == "":
                line = synoantonym_string(line, excluded_words, synonym)
                output_file.write(line + "\n")
            else:
                output_file.write("\n")


__version__ = '1.0'
__author__ = 'Oren Anderson'


# Create a subclass of QMainWindow to setup the calculator's GUI
class ThesaurusPlusUi(QMainWindow):
    def __init__(self):
        super().__init__()

        # Set some main window's properties
        self.setWindowTitle('Thesaurus Plus')
        self.setWindowIcon(QtGui.QIcon('C:\\Users\oren\Documents\\0_Projects\\2021\ThesaurusPlus\src\main\icons\Phrog.ico'))
        self.setMinimumSize(700, 400)

        # Initialize re-running dialog box popup
        self.dialog_box = PopUpBoxReRun(self)
        self.dialog_config = PopUpBoxConfig(self)

        # Set Central Widget
        self._central_widget = QWidget(self)
        self.setCentralWidget(self._central_widget)

        # Set the general layout
        self._general_layout = QVBoxLayout()
        self._central_widget.setLayout(self._general_layout)

        # Create menu and the content/text fields
        self._create_menu()
        self._create_content()

        self.input_text.setFocus()

    def _create_menu(self):
        self.top_menu_layout = QHBoxLayout()
        self._general_layout.addLayout(self.top_menu_layout)

        self.run = QPushButton('Run')
        self.top_menu_layout.addWidget(self.run)
        self.re_running = QPushButton('Re-Running')
        self.top_menu_layout.addWidget(self.re_running)
        self.import_text = QPushButton('Import')
        self.top_menu_layout.addWidget(self.import_text)
        self.export_text = QPushButton('Export')
        self.top_menu_layout.addWidget(self.export_text)
        self.configure = QPushButton('Configure')
        self.top_menu_layout.addWidget(self.configure)

    def _create_content(self):
        self.primary_layout = QVBoxLayout()
        self._general_layout.addLayout(self.primary_layout)

        self.primary_layout.addWidget(QLabel('Input Text'))
        self.input_text = QPlainTextEdit('Type text here or use the import button to import larger files')
        self.primary_layout.addWidget(self.input_text)

        self.primary_layout.addWidget(QLabel('Output Text'))
        self.output_text = QPlainTextEdit()
        self.primary_layout.addWidget(self.output_text)

    def set_input_text(self, text):
        self.input_text.setPlainText(text)
        self.input_text.setFocus()

    def get_input_text(self):
        return self.input_text.toPlainText()

    def set_output_text(self, text):
        self.output_text.setPlainText(text)
        self.output_text.setFocus()

    def get_output_text(self):
        return self.output_text.toPlainText()


class PopUpBoxReRun(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent, QtCore.Qt.WindowCloseButtonHint)
        self.setWindowTitle('Re-Running Options')
        self.setWindowIcon(QtGui.QIcon('C:\\Users\oren\Documents\\0_Projects\\2021\ThesaurusPlus\src\main\icons\Phrog.ico'))
        self.setMinimumWidth(200)
        dlg_layout = QVBoxLayout()
        form_layout = QFormLayout()

        # re-run inputs
        self.onlyInt = QIntValidator(1, 100, self)
        self.number_of_runs = QLineEdit('5')
        self.number_of_runs.setValidator(self.onlyInt)

        self.detailed_output = QCheckBox()

        form_layout.addRow('Number of runs:', self.number_of_runs)
        form_layout.addRow('Detailed Output:', self.detailed_output)
        dlg_layout.addLayout(form_layout)

        self.btns = QDialogButtonBox()
        self.btns.setStandardButtons(QDialogButtonBox.Cancel | QDialogButtonBox.Ok)
        dlg_layout.addWidget(self.btns)

        self.setLayout(dlg_layout)

    def closeEvent(self, *args, **kwargs):
        self.parent().setEnabled(True)
        self.close()

    def accepted(self):
        self.close()

    def rejected(self):
        self.close()

    def display_dialog(self):
        self.show()


class PopUpBoxConfig(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent, QtCore.Qt.WindowCloseButtonHint)
        self.setWindowTitle('Configuration')
        self.setWindowIcon(QtGui.QIcon('C:\\Users\oren\Documents\\0_Projects\\2021\ThesaurusPlus\src\main\icons\Phrog.ico'))
        self.setMinimumWidth(200)

        config_dir = str(Path(__file__).parent.resolve())

        main_layout = QVBoxLayout()
        main_layout.addWidget(QLabel('Configuration File In: ' + config_dir))
        main_layout.addWidget(QLabel('Ignored Words (Separate With Spaces)'))
        self.input_text = QPlainTextEdit()
        main_layout.addWidget(self.input_text)

        form_layout = QFormLayout()
        self.antonym_mode = QCheckBox()
        form_layout.addRow('Antonym Mode:', self.antonym_mode)
        main_layout.addLayout(form_layout)

        self.btns = QDialogButtonBox()
        self.btns.setStandardButtons(QDialogButtonBox.Cancel | QDialogButtonBox.Ok | QDialogButtonBox.Save)
        main_layout.addWidget(self.btns)

        self.setLayout(main_layout)

    def closeEvent(self, *args, **kwargs):
        self.parent().setEnabled(True)
        self.close()

    def accepted(self):
        self.close()

    def rejected(self):
        self.close()

    def display_dialog(self):
        self.show()


# Create a Controller class to connect the UI
class Controller:
    def __init__(self, view):
        self._view = view
        # Connect signals and slots
        self._connect_signals()
        config = load_settings()
        self._ignored_words = config[0]
        self._synonym_mode = config[1]
        self._view.dialog_config.input_text.setPlainText(self._ignored_words)

    def _import_text(self):
        # open file dialog box, navigation starting at home directory
        home_dir = str(Path.home())
        file_name = QFileDialog.getOpenFileName(caption='Open Text File', directory=home_dir,
                                                filter="txt files (*.txt *.docx)")
        file_name = file_name[0]

        extension = os.path.splitext(file_name)[1]

        if "txt" in extension:
            if file_name:
                with open(file_name, "r") as fh:
                    import_text = fh.read()
                    self._view.set_input_text(import_text)
        elif "docx" in extension:
            document = docx.Document(file_name)

            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            full_text = '\n'.join(full_text)

            self._view.set_input_text(full_text)

        self._view.input_text.setFocus()

    def _export_text(self):
        # export output text navigation starting at home directory
        home_dir = str(Path.home())
        file_name = QFileDialog.getSaveFileName(caption='Save Text File', directory=home_dir,
                                                filter="txt files (*.txt *.docx)")
        file_name = file_name[0]
        extension = os.path.splitext(file_name)[1]
        if "txt" in extension:
            if file_name:
                with open(file_name, "w+") as fh:
                    fh.write(self._view.get_output_text())
        elif "docx" in extension:
            document = docx.Document()
            document.add_paragraph(self._view.get_output_text())
            document.save(file_name)

        self._view.input_text.setFocus()

    def _run_input(self):
        text_input = self._view.get_input_text()
        output = synoantonym_string(text_input, self._ignored_words, self._synonym_mode)
        self._view.set_output_text(output)

    def _on_re_run_accept(self):
        number_of_runs = int(self._view.dialog_box.number_of_runs.text())
        show_intermediate = bool(self._view.dialog_box.detailed_output.checkState())
        self._view.setEnabled(True)
        self._view.dialog_box.setEnabled(True)
        self._view.dialog_box.accepted()
        self._re_run_action(number_of_runs, show_intermediate)
        self._view.input_text.setFocus()

    def _on_re_run_reject(self):
        self._view.setEnabled(True)
        self._view.dialog_box.setEnabled(True)
        self._view.dialog_box.rejected()

    def _on_re_run_start(self):
        self._view.setEnabled(False)
        self._view.dialog_box.setEnabled(True)
        self._view.dialog_box.display_dialog()

    def _re_run_action(self, number_of_runs, show_intermediate):
        total = ""
        output = self._view.get_input_text()
        for i in range(number_of_runs):
            output = synoantonym_string(output, self._ignored_words, self._synonym_mode)
            if show_intermediate:
                if i == number_of_runs - 1:
                    total = total + str(i + 1) + ": " + output
                else:
                    total = total + str(i + 1) + ": " + output + "\n"
        if show_intermediate:
            self._view.set_output_text(total)
        else:
            self._view.set_output_text(output)

    def _on_config_accept(self):
        self._ignored_words = self._view.dialog_config.input_text.toPlainText()
        self._synonym_mode = not bool(self._view.dialog_config.antonym_mode.checkState())
        self._view.setEnabled(True)
        self._view.dialog_config.setEnabled(True)
        self._view.dialog_config.accepted()

    def _on_config_reject(self):
        self._view.setEnabled(True)
        self._view.dialog_config.setEnabled(True)
        self._view.dialog_config.rejected()

    def _on_config_start(self):
        self._view.dialog_config.input_text.setPlainText(self._ignored_words)
        self._view.dialog_config.antonym_mode.setChecked(not self._synonym_mode)
        self._view.setEnabled(False)
        self._view.dialog_config.setEnabled(True)
        self._view.dialog_config.display_dialog()

    def _handle_dialog_buttons(self, button):
        button = button.text()
        if button in "Save" or button in "OK":
            self._on_config_accept()

        if button in "Save":
            save_settings(self._ignored_words, self._synonym_mode)

        if button in "Cancel":
            self._on_config_reject()

        self._view.input_text.setFocus()

    def _connect_signals(self):
        # self.dialog_config
        self._view.run.clicked.connect(partial(self._run_input))
        self._view.re_running.clicked.connect(partial(self._on_re_run_start))
        self._view.import_text.clicked.connect(partial(self._import_text))
        self._view.export_text.clicked.connect(partial(self._export_text))
        self._view.configure.clicked.connect(partial(self._on_config_start))

        self._view.dialog_box.btns.accepted.connect(partial(self._on_re_run_accept))
        self._view.dialog_box.btns.rejected.connect(partial(self._on_re_run_reject))

        self._view.dialog_config.btns.clicked.connect(partial(self._handle_dialog_buttons))


def load_settings():
    if path.exists("config.txt"):
        with open("config.txt", "r") as fh:
            import_text = fh.read()
            import_text = import_text.split("\n")

            try:
                ignored_words = import_text[0][import_text[0].find("=") + 1:].strip()
                antonym_mode = import_text[1][import_text[1].find("=") + 1:].strip()
                if antonym_mode == 'True':
                    antonym_mode = False
                elif antonym_mode == 'False':
                    antonym_mode = True
            except ValueError:
                print("Config File Not Set Up Properly")

            return ignored_words, antonym_mode
    else:
        with open("config.txt", "w+") as fh:
            fh.write("ignored words = a an the i it as its no in test\nantonym mode = False")
            return "a an the i it as its no in", True


def save_settings(ignored_words, antonym_mode):
    ignored_words = "ignored words = " + ignored_words
    antonym_mode = "antonym mode = " + str(not antonym_mode)
    output = ignored_words + "\n" + antonym_mode

    with open("config.txt", "w") as fh:
        fh.write(output)


# Client code
def main():
    """Main function."""
    # Create an instance of `QApplication`
    appctxt = ApplicationContext()  # 1. Instantiate ApplicationContext

    # Load qss (qt equivalent of css)
    sshFile = "C:\\Users\oren\Documents\\0_Projects\\2021\ThesaurusPlus\src\main\python\customLooks.qss"
    with open(sshFile, "r") as fh:
        appctxt.app.setStyleSheet(fh.read())

    # Show the calculator's GUI
    view = ThesaurusPlusUi()
    view.show()

    # Create instances of the model and the controller
    Controller(view)

    # Download or update language toolkit
    download()

    # Execute calculator's main loop
    exit_code = appctxt.app.exec()  # 2. Invoke appctxt.app.exec()
    sys.exit(exit_code)


if __name__ == '__main__':
    main()
